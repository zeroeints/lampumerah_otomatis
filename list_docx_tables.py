import docx

def main():
    doc = docx.Document("2215354055_Mohammad Filla Firdaus bab 1^J2^03.docx")
    print(f"Total Tables in doc: {len(doc.tables)}")
    
    # Let's inspect the paragraphs around each table to find its caption
    # Usually the caption is in a paragraph just before or after the table.
    # Let's print table index, and the first few rows of each table.
    for i, table in enumerate(doc.tables):
        print(f"\n====================================")
        print(f"TABLE {i} (Rows: {len(table.rows)}, Cols: {len(table.columns)})")
        print(f"====================================")
        
        # Let's find text in the first row
        try:
            header = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
            print("Header:", header[:5])
        except Exception as e:
            print("Error reading header:", e)
            
        # Print first row of data if available
        if len(table.rows) > 1:
            try:
                row1 = [cell.text.strip().replace("\n", " ") for cell in table.rows[1].cells]
                print("Row 1 :", row1[:5])
            except Exception as e:
                print("Error reading row 1:", e)

if __name__ == "__main__":
    main()

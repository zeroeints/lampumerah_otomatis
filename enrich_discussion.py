import docx
from docx.shared import Pt, Cm
import os

def main():
    doc_path = "2215354055_Mohammad Filla Firdaus bab 1^J2^03.docx"
    doc = docx.Document(doc_path)
    print(f"Loaded document with {len(doc.paragraphs)} paragraphs.")
    
    # We will locate paragraph indices by matching their start text.
    # Paragraph 352: "Sebaliknya, kelas Sepeda Motor (*motorcycle*) mencatatkan recall yang sangat rendah yaitu 26,10%..."
    # Paragraph 355: "Penerapan pembobotan smp MKJI 1997 ini berhasil menciptakan keadilan keputusan durasi sinyal hijau..."
    # Paragraph 361: "Selain itu, pembatasan durasi merah maksimum (Max Red-Light Duration) yang tetap terjaga stabil pada kapasitas 38,0 detik..."
    
    idx_352 = None
    idx_355 = None
    idx_361 = None
    
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("Sebaliknya, kelas Sepeda Motor") and "26,10%" in text:
            idx_352 = idx
        elif text.startswith("Penerapan pembobotan smp MKJI 1997 ini berhasil menciptakan keadilan"):
            idx_355 = idx
        elif text.startswith("Selain itu, pembatasan durasi merah maksimum") and "38,0 detik" in text:
            idx_361 = idx
            
    print(f"Found targets: 352={idx_352}, 355={idx_355}, 361={idx_361}")
    
    if idx_352 is not None:
        p = doc.paragraphs[idx_352]
        p.text = (
            "Sebaliknya, kelas Sepeda Motor (motorcycle) mencatatkan recall yang sangat rendah yaitu 26,10% meskipun memiliki "
            "presisi sebesar 73,30%. Hal ini disebabkan oleh fenomena oklusi fisik (spatial occlusion) dan distorsi spasial, di mana "
            "sepeda motor sering kali berjalan berhimpitan sangat rapat di lajur pendekat, menyebabkan model mendeteksi kerumunan "
            "motor sebagai satu kesatuan objek atau melewatkan motor di bagian belakang. Selain itu, terdapat bias minor pada kategori "
            "kendaraan berat (Heavy Vehicle), di mana Bus (bus) terdeteksi sangat baik dengan mAP@50 sebesar 92,90%, sedangkan "
            "Truk (truck) hanya mencapai 66,40%. Bias ini disebabkan oleh kemiripan geometri visual kotak persegi panjang dari "
            "sudut pandang CCTV atas (top-down view), sehingga truk berukuran sedang sering kali disalahklasifikasikan sebagai bus, "
            "dan sebaliknya. Rendahnya nilai recall sepeda motor ini merupakan bentuk galat klasifikasi riil yang disebabkan oleh "
            "faktor lingkungan dan kerumunan objek. Pada sistem kontrol lalu lintas konvensional berbasis persamaan matematis linear, "
            "galat deteksi ini akan langsung memicu pemotongan waktu hijau secara drastis (under-allocation) karena jumlah kendaraan "
            "terhitung lebih sedikit dari kondisi aslinya. Namun, keunggulan utama dari perancangan model terintegrasi ini adalah "
            "penerapan Logika Fuzzy Mamdani sebagai pengendali keputusan yang bertindak sebagai filter toleransi kesalahan (error "
            "tolerance filter). Logika fuzzy tidak bergantung pada persamaan linear kaku, melainkan menggunakan himpunan fuzzy yang "
            "tumpang tindih (overlapping membership functions) untuk mengklasifikasikan tingkat kepadatan lajur, sehingga mampu "
            "meredam dampak galat deteksi YOLOv11s dan tetap menghasilkan keputusan durasi hijau yang stabil."
        )
        p.alignment = 3
        p.paragraph_format.first_line_indent = Cm(1.0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
    if idx_355 is not None:
        p = doc.paragraphs[idx_355]
        p.text = (
            "Penerapan pembobotan smp MKJI 1997 ini berhasil menciptakan keadilan keputusan durasi sinyal hijau (phase fairness). "
            "Tanpa bobot smp (hanya mengandalkan jumlah kendaraan mentah), lajur yang dipenuhi oleh 30 sepeda motor akan mendapat "
            "durasi hijau maksimal yang sangat panjang secara tidak adil, padahal ruang jalan yang terpakai sangat kecil. Dengan "
            "bobot smp, 30 motor tersebut dikonversi menjadi 6 smp, setara dengan ruang jalan 6 mobil, sehingga sistem mengalokasikan "
            "durasi hijau yang lebih efisien dan proporsional. Sebaliknya, lajur dengan volume kendaraan berat yang tinggi (bus/truk) "
            "mendapat bobot lebih besar (1.3 smp per unit) karena membutuhkan waktu pembebasan simpang (discharge time) yang lebih "
            "lama akibat akselesari awal yang lambat, sehingga mencegah kemacetan terkunci di tengah persimpangan. "
            "Pembobotan smp MKJI 1997 ini juga secara tidak langsung memperkuat ketahanan sistem terhadap galat deteksi YOLO. Karena sepeda "
            "motor yang rentan mengalami oklusi spasial hanya berbobot 0,2 smp, kesalahan deteksi beberapa unit sepeda motor tidak akan "
            "mengubah total nilai kepadatan (smp) secara ekstrem pada semesta pembicaraan fuzzy. Sebagai contoh, selisih deteksi akibat "
            "oklusi sebanyak 10 motor hanya berdampak sebesar 2,0 smp. Selisih kecil ini dapat dengan mudah diredam oleh fungsi "
            "keanggotaan fuzzy karena nilai kepadatan aktual dan terdeteksi tetap berada pada derajat keanggotaan yang dominan di "
            "himpunan yang sama (misalnya, kelas 'Sedikit' atau 'Sedang'). Selain itu, penggunaan Panjang Antrean (queue length) "
            "sebagai input kedua dari data internal simulator bertindak sebagai parameter redundansi. Jika YOLOv11s meleset mendeteksi "
            "volume kendaraan karena oklusi ekstrem, input panjang antrean yang presisi akan tetap memicu aturan fuzzy untuk memberikan "
            "durasi hijau yang memadai, sehingga keputusan durasi lampu hijau tetap responsif dan adil."
        )
        p.alignment = 3
        p.paragraph_format.first_line_indent = Cm(1.0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
    if idx_361 is not None:
        p = doc.paragraphs[idx_361]
        p.text = (
            "Selain itu, pembatasan durasi merah maksimum (Max Red-Light Duration) yang tetap terjaga stabil pada kapasitas 38,0 detik "
            "bertindak sebagai batas konstrain (boundary constraint) yang aman. Batasan ini menjamin lajur pendekat lain yang sepi "
            "tidak akan mengalami penundaan tanpa batas (starvation) akibat lajur pendekat utama terus menerus memperpanjang fase hijaunya. "
            "Hasil akhir menunjukkan perbaikan indeks kepadatan persimpangan (Density Index) secara total sebesar 60,98% lebih lancar "
            "(turun dari kepadatan 20,5% ke 8,0%), membuktikan sistem adaptif ini bekerja secara optimal untuk mengurai kemacetan perkotaan. "
            "Batas konstrain ini juga bertindak sebagai jaring pengaman (safety net) jika terjadi kegagalan deteksi total pada kamera "
            "CCTV (misalnya akibat lensa terhalang debu atau cuaca ekstrem). Dengan jaring pengaman ini, fase lampu hijau minimal 10 "
            "detik dan fase lampu merah maksimal 38 detik tetap dipatuhi, mencegah persimpangan terkunci total (deadlock) akibat ketiadaan "
            "input deteksi. Integrasi sirkuit tertutup (closed-loop feedback) ini secara empiris membuktikan bahwa kombinasi detektor "
            "YOLOv11s dan pengontrol logika fuzzy mampu menoleransi galat masukan (input errors) jauh lebih baik dibandingkan sistem "
            "kendali lampu lalu lintas berbasis persamaan matematis statis konvensional."
        )
        p.alignment = 3
        p.paragraph_format.first_line_indent = Cm(1.0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
    doc.save(doc_path)
    print("[OK] Document enriched with academic fuzzy error-tolerance discussion successfully!")

if __name__ == "__main__":
    main()

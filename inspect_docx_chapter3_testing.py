import docx

def find_table_by_title(doc, title_keywords):
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if any(kw in text.lower() for kw in title_keywords):
            print(f"\nFound paragraph [{i}]: {text}")
            # Check if there is a table nearby (usually next element)
            # Let's check paragraphs around it
            # Let's find which table is closest to this paragraph in the document
            
def main():
    doc = docx.Document("2215354055_Mohammad Filla Firdaus bab 1^J2^03.docx")
    
    # We want to print the text of Tables in Chapter 3
    # Let's find tables that correspond to table captions
    print("=== SEARCHING TABLES IN CHAPTER 3 ===")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "Tabel 3." in text or "Tabel 4." in text:
            print(f"[{i}] {text}")
            
    # Let's inspect doc.tables content for index 7, 8, 9, 10, 11
    # In list_docx_tables we saw there are 17 tables total.
    # Let's print tables 7 to 13 to identify them.
    for idx in range(7, 14):
        if idx < len(doc.tables):
            table = doc.tables[idx]
            print(f"\nTable {idx}: rows={len(table.rows)}, cols={len(table.columns)}")
            try:
                row0 = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
                print(f"  Header: {row0[:4]}")
                if len(table.rows) > 1:
                    row1 = [cell.text.strip().replace("\n", " ") for cell in table.rows[1].cells]
                    print(f"  Row 1 : {row1[:4]}")
            except Exception as e:
                print("  Error:", e)

if __name__ == "__main__":
    main()

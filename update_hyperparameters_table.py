import docx
from docx.shared import Pt

def main():
    doc_path = "2215354055_Mohammad Filla Firdaus bab 1^J2^03.docx"
    doc = docx.Document(doc_path)
    print(f"Loaded document with {len(doc.tables)} tables.")
    
    # Table 9 is the hyperparameter table
    t = doc.tables[9]
    print("Modifying Table 9 (Tabel 3.7)...")
    
    # New row values
    actual_rows = [
        ["1.", "Epoch", "100", "Jumlah iterasi penuh pelatihan pada seluruh Dataset untuk mencegah underfitting atau overfitting."],
        ["2.", "Batch Size", "8", "Jumlah citra yang diproses dalam satu kali langkah pembaruan bobot, disesuaikan dengan kapasitas VRAM GPU Tesla T4."],
        ["3.", "Learning Rate", "0.0005", "Ukuran langkah adaptasi model saat memperbarui bobot (learning rate awal)."],
        ["4.", "Optimizer", "AdamW", "Algoritma yang digunakan untuk mengoptimalkan pembaruan bobot selama proses pelatihan."],
        ["5.", "Image Size", "960 x 960", "Resolusi input gambar yang diproses oleh model untuk menjaga keseimbangan antara akurasi dan kecepatan."]
    ]
    
    for r_idx, data in enumerate(actual_rows):
        row = t.rows[r_idx + 1] # Skip header
        for c_idx, val in enumerate(data):
            cell = row.cells[c_idx]
            cell.text = val
            # Re-apply font formatting (Times New Roman 11pt, normal style)
            p = cell.paragraphs[0]
            p.alignment = 1 if c_idx in [0, 2] else 0  # Center align for No and Value
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = False
                
    doc.save(doc_path)
    print("[OK] Hyperparameters table updated successfully in Word document!")

if __name__ == "__main__":
    main()

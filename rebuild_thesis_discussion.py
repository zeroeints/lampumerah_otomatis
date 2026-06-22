import docx
from docx.shared import Inches, Pt, Cm
import os

def insert_body_paragraph_before(paragraph, text):
    p = paragraph.insert_paragraph_before()
    p.alignment = 3  # Justified
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    
    parts = text.split("**")
    for idx, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if idx % 2 == 1:
            run.bold = True
    return p

def insert_heading_before(paragraph, text, level=3):
    p = paragraph.insert_paragraph_before()
    p.alignment = 0  # Left
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p

def insert_caption_before(paragraph, text):
    p = paragraph.insert_paragraph_before()
    p.alignment = 1  # Centered
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = True
    return p

def main():
    doc_path = "2215354055_Mohammad Filla Firdaus bab 1^J2^03.docx"
    doc = docx.Document(doc_path)
    print(f"Loaded document with {len(doc.paragraphs)} paragraphs.")
    
    # 1. Move performance benchmarks to 4.2.3
    # Currently, paragraphs 347 to 355 contain:
    # 347: Intro paragraph ("Selain parameter...")
    # 348: YOLO latency bullet
    # 349: Fuzzy latency bullet
    # 350: RAM memory bullet
    # 351: TraCI stability bullet
    # 352: Blank/Paragraph spacing
    # 353: Image 4.6 (Computational Latency)
    # 354: Blank/Paragraph spacing
    # 355: Image 4.7 (RAM Memory)
    #
    # Wait, we want these to be located BEFORE paragraph 346 ("4.3 Analisis Hasil Pengujian").
    # If we move them before 346, they will naturally be placed under 4.2.3 (after Table 4.6).
    # Let's perform this relocation by copying the text/elements and deleting the old ones.
    
    # Let's verify the paragraphs at these indices
    p346 = doc.paragraphs[346]
    print(f"p346 text: '{p346.text}'")
    
    # Let's store the content of paragraphs 347 to 355
    moved_paras = []
    for idx in range(347, 356):
        moved_paras.append(doc.paragraphs[idx])
        
    print(f"[+] Relocating {len(moved_paras)} paragraphs of performance benchmarks...")
    
    # We will insert them BEFORE p346 (which is heading 4.3).
    # Since we insert before p346, they will be placed after Table 4.6 and under 4.2.3.
    # To do this cleanly, we can recreate the paragraphs before p346 and then delete the old ones.
    
    # We will recreate the paragraphs before p346:
    # 1. Intro
    p_intro = p346.insert_paragraph_before()
    p_intro.alignment = 3
    p_intro.paragraph_format.first_line_indent = Cm(1.0)
    p_intro.paragraph_format.line_spacing = 1.5
    p_intro.paragraph_format.space_before = Pt(6)
    p_intro.paragraph_format.space_after = Pt(6)
    r = p_intro.add_run(doc.paragraphs[347].text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    
    # 2. Bullets
    for b_idx in [348, 349, 350, 351]:
        p_b = p346.insert_paragraph_before()
        p_b.alignment = 3
        p_b.paragraph_format.first_line_indent = Cm(1.0)
        p_b.paragraph_format.line_spacing = 1.5
        p_b.paragraph_format.space_before = Pt(0)
        p_b.paragraph_format.space_after = Pt(4)
        
        # Split bold prefix
        orig_text = doc.paragraphs[b_idx].text
        parts = orig_text.split(":")
        if len(parts) >= 2:
            prefix = parts[0] + ":"
            suffix = ":".join(parts[1:])
            r1 = p_b.add_run(prefix)
            r1.bold = True
            r1.font.name = 'Times New Roman'
            r1.font.size = Pt(12)
            r2 = p_b.add_run(suffix)
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(12)
        else:
            r = p_b.add_run(orig_text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            
    # 3. Image 4.6
    p_img1 = p346.insert_paragraph_before()
    p_img1.alignment = 1
    p_img1.paragraph_format.space_before = Pt(12)
    p_img1.paragraph_format.space_after = Pt(4)
    run_img1 = p_img1.add_run()
    run_img1.add_picture("output/latensi_komputasi_sistem.png", width=Inches(5.5))
    
    p_cap1 = p346.insert_paragraph_before()
    p_cap1.alignment = 1
    p_cap1.paragraph_format.space_before = Pt(4)
    p_cap1.paragraph_format.space_after = Pt(12)
    run_cap1 = p_cap1.add_run("Gambar 4.6 Grafik Perbandingan Latensi Pemrosesan Komputasi Sistem")
    run_cap1.bold = True
    run_cap1.font.name = 'Times New Roman'
    run_cap1.font.size = Pt(11)
    
    # 4. Image 4.7
    p_img2 = p346.insert_paragraph_before()
    p_img2.alignment = 1
    p_img2.paragraph_format.space_before = Pt(12)
    p_img2.paragraph_format.space_after = Pt(4)
    run_img2 = p_img2.add_run()
    run_img2.add_picture("output/penggunaan_memori_ram.png", width=Inches(5.5))
    
    p_cap2 = p346.insert_paragraph_before()
    p_cap2.alignment = 1
    p_cap2.paragraph_format.space_before = Pt(4)
    p_cap2.paragraph_format.space_after = Pt(12)
    run_cap2 = p_cap2.add_run("Gambar 4.7 Grafik Distribusi Alokasi Memori RAM Saat Sistem Beroperasi")
    run_cap2.bold = True
    run_cap2.font.name = 'Times New Roman'
    run_cap2.font.size = Pt(11)
    
    # Now we save the document, reload it, and delete the OLD paragraphs (which are now shifted down).
    # Since we inserted 9 paragraphs before p346, the old paragraphs that were at 347-355 are now at indices 347+9 = 356 to 364.
    # Actually, it's safer to just delete them by matching their text to avoid index offset bugs!
    doc.save(doc_path)
    
    # Reload and delete
    doc = docx.Document(doc_path)
    # We delete paragraphs whose text matches the old moved paragraphs.
    # Note: the images are in run.add_picture and their paragraphs don't have text, but they have captions.
    # We can delete paragraphs by looking for captions 'Gambar 4.6' and 'Gambar 4.7' that appear AFTER paragraph 350.
    # Let's write a clean deletion loop.
    to_delete = []
    caption_count = 0
    for idx, p in enumerate(doc.paragraphs):
        # The OLD ones will be located after the 4.3 heading which is now shifted down.
        # Let's find the second occurrence of 'Gambar 4.6' or 'Gambar 4.7'
        if "Gambar 4.6" in p.text or "Gambar 4.7" in p.text:
            caption_count += 1
            if caption_count > 2: # Keep the first two (which are the newly inserted ones), delete the later ones.
                to_delete.append(idx)
                # Also delete the image paragraph just before the caption
                to_delete.append(idx - 1)
        elif idx > 356 and any(prefix in p.text for prefix in ["Selain parameter efisiensi rekayasa", "Latensi Inferensi YOLOv11s:", "Latensi Keputusan Pengontrol", "Alokasi Kapasitas Memori (RAM):", "Stabilitas Transmisi Protokol"]):
            to_delete.append(idx)
            
    # Delete in reverse order
    for idx in sorted(list(set(to_delete)), reverse=True):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)
        
    doc.save(doc_path)
    print("[+] Successfully relocated benchmarks to 4.2.3 and deleted duplicates.")
    
    # -------------------------------------------------------------
    # 2. Rename Heading 4.3 and Rebuild Section 4.3 Content
    # -------------------------------------------------------------
    doc = docx.Document(doc_path)
    
    # Find heading 4.3
    heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("4.3"):
            heading_idx = i
            break
            
    if heading_idx is None:
        print("[!] Heading 4.3 not found!")
        return
        
    print(f"[+] Found heading 4.3 at paragraph index: {heading_idx}")
    doc.paragraphs[heading_idx].text = "4.3 Pembahasan Akademis dan Analisis Komparatif"
    for run in doc.paragraphs[heading_idx].runs:
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
    # Find BAB V to delete everything in between
    bab5_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("BAB V"):
            bab5_idx = i
            break
            
    if bab5_idx is None:
        print("[!] BAB V not found!")
        return
        
    print(f"[+] Found BAB V at paragraph index: {bab5_idx}")
    
    # Delete paragraphs between heading_idx and bab5_idx
    # Let's collect them first
    paras_to_delete = list(range(heading_idx + 1, bab5_idx))
    print(f"[+] Deleting {len(paras_to_delete)} old paragraphs under section 4.3...")
    for idx in sorted(paras_to_delete, reverse=True):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)
        
    doc.save(doc_path)
    
    # Reload and insert the new academic section
    doc = docx.Document(doc_path)
    
    # Find BAB V again
    bab5_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("BAB V"):
            bab5_idx = i
            break
            
    target_p = doc.paragraphs[bab5_idx]
    
    # --- INSERT SUBSECTION 4.3.1 ---
    insert_heading_before(target_p, "4.3.1 Analisis Validitas Perancangan Integrasi Visi Komputer dan Logika Fuzzy (Menjawab Rumusan Masalah 1)", level=3)
    
    insert_body_paragraph_before(target_p,
        "Pengujian fungsional sistem membuktikan bahwa perancangan integrasi antara subsistem visi komputer (YOLOv11s) "
        "dan mesin inferensi logika fuzzy yang dihubungkan secara asinkron ke simulator SUMO melalui protokol TraCI "
        "memiliki validitas operasional yang sangat tinggi. Koordinasi asinkron antara dashboard **Streamlit** dan "
        "sirkuit kendali Python backend berhasil mengatasi kendala desinkronisasi waktu jabat tangan (**TraCI handshake**) "
        "yang sering terjadi pada sistem kendali lalu lintas berbasis socket TCP/IP."
    )
    
    insert_body_paragraph_before(target_p,
        "Validitas ini didukung oleh hasil pengukuran latensi komputasi, di mana latensi inferensi YOLOv11s sebesar **9,7 ms** "
        "pada GPU Tesla T4 dan latensi keputusan fuzzy sebesar **4,3 ms** berada jauh di bawah batas konstrain waktu nyata "
        "(**real-time constraint < 100 ms**). Kecepatan komputasi ini menjamin bahwa setiap data volume kendaraan terdeteksi "
        "dapat dikonversi secara instan menjadi keputusan sinyal lampu lalu lintas baru pada langkah simulasi berikutnya tanpa "
        "menimbulkan jeda (*lag*) komunikasi socket, sehingga membuktikan keandalan integrasi sistem tertutup (*closed-loop*)."
    )
    
    # --- INSERT SUBSECTION 4.3.2 ---
    insert_heading_before(target_p, "4.3.2 Analisis Karakteristik Performa Deteksi Multi-Kelas YOLOv11 (Menjawab Rumusan Masalah 2)", level=3)
    
    insert_body_paragraph_before(target_p,
        "Penerapan model deteksi objek YOLOv11 Small (YOLOv11s) menunjukkan sensitivitas dan akurasi klasifikasi yang tinggi "
        "pada empat kelas target kendaraan, dengan skor **mAP@50 mencapai 69,23%** secara keseluruhan. Namun, analisis data "
        "per-kelas menunjukkan karakteristik performa yang berbeda secara signifikan. Kelas Mobil (*car*) mencatatkan akurasi "
        "presisi tertinggi sebesar **89,60%** dan recall **83,30%** karena bentuk geometris mobil yang relatif konsisten dan "
        "jelas pada sudut pandang CCTV atas persimpangan."
    )
    
    insert_body_paragraph_before(target_p,
        "Sebaliknya, kelas Sepeda Motor (*motorcycle*) mencatatkan recall yang sangat rendah yaitu **26,10%** meskipun memiliki presisi "
        "sebesar **73,30%**. Hal ini disebabkan oleh fenomena oklusi fisik (*spatial occlusion*) dan distorsi spasial, di mana "
        "sepeda motor sering kali berjalan berhimpitan sangat rapat di lajur pendekat, menyebabkan model mendeteksi kerumunan "
        "motor sebagai satu kesatuan objek atau melewatkan motor di bagian belakang. Selain itu, terdapat bias minor pada kategori "
        "kendaraan berat (*Heavy Vehicle*), di mana Bus (*bus*) terdeteksi sangat baik dengan **mAP@50 sebesar 92,90%**, sedangkan "
        "Truk (*truck*) hanya mencapai **66,40%**. Bias ini disebabkan oleh kemiripan geometri visual kotak persegi panjang dari "
        "sudut pandang CCTV atas (*top-down view*), sehingga truk berukuran sedang sering kali disalahklasifikasikan sebagai bus, "
        "dan sebaliknya."
    )
    
    # --- INSERT SUBSECTION 4.3.3 ---
    insert_heading_before(target_p, "4.3.3 Analisis Responsivitas Aturan Fuzzy Berbasis Pembobotan smp MKJI 1997 (Menjawab Rumusan Masalah 3)", level=3)
    
    insert_body_paragraph_before(target_p,
        "Mesin inferensi logika fuzzy yang dirancang menggunakan 2 variabel input (Kepadatan Kendaraan dan Panjang Antrean) "
        "serta 9 aturan keputusan (*3x3 grid rule base*) terbukti responsif dalam menentukan durasi lampu hijau secara dinamis. "
        "Signifikansi pengujian ini terletak pada integrasi koefisien **Ekuivalensi Mobil Penumpang (EMP)** berdasarkan panduan "
        "**Manual Kapasitas Jalan Indonesia (MKJI) 1997** (Sepeda Motor: 0.2, Mobil: 1.0, Bus/Truk: 1.3) untuk mengonversi jumlah "
        "kendaraan mentah menjadi volume terbobot Satuan Mobil Penumpang (smp)."
    )
    
    insert_body_paragraph_before(target_p,
        "Penerapan pembobotan smp MKJI 1997 ini berhasil menciptakan keadilan keputusan durasi sinyal hijau (*phase fairness*). "
        "Tanpa bobot smp (hanya mengandalkan jumlah kendaraan mentah), lajur yang dipenuhi oleh 30 sepeda motor akan mendapat "
        "durasi hijau maksimal yang sangat panjang, padahal ruang jalan yang terpakai sangat kecil. Dengan bobot smp, 30 motor "
        "tersebut dikonversi menjadi **6 smp**, setara dengan ruang jalan 6 mobil, sehingga sistem mengalokasikan durasi hijau yang "
        "lebih efisien dan proporsional. Sebaliknya, lajur dengan volume kendaraan berat yang tinggi (bus/truk) mendapat bobot "
        "lebih besar (1.3 smp per unit) karena membutuhkan waktu pembebasan simpang (*discharge time*) yang lebih lama akibat akselerasi "
        "awal yang lambat, sehingga mencegah kemacetan terkunci di tengah persimpangan."
    )
    
    # --- INSERT SUBSECTION 4.3.4 ---
    insert_heading_before(target_p, "4.3.4 Analisis Efektivitas Mekanisme Green Extension dan Green Cut terhadap Kinerja Lalu Lintas (Menjawab Rumusan Masalah 4)", level=3)
    
    insert_body_paragraph_before(target_p,
        "Keberhasilan rekayasa lalu lintas makro pada persimpangan terbukti sangat signifikan melalui penerapan kontroler adaptif "
        "fuzzy. Sistem ini sukses memotong waktu tunggu rata-rata kendaraan (**Avg Waiting Time**) sebesar **62,67% lebih cepat** "
        "(turun dari 4,308 detik menjadi 1,608 detik) dan mengurangi rata-rata panjang antrean (**Avg Queue Length**) sebesar "
        "**61,01% lebih pendek** (turun dari 16,399 unit menjadi 6,394 unit). Penurunan antrean secara kontinu selama 1.800 detik "
        "langkah simulasi divisualisasikan secara komparatif pada Gambar 4.8."
    )
    
    # Insert Gambar 4.8
    p_img3 = target_p.insert_paragraph_before()
    p_img3.alignment = 1
    p_img3.paragraph_format.space_before = Pt(12)
    p_img3.paragraph_format.space_after = Pt(4)
    run_img3 = p_img3.add_run()
    run_img3.add_picture("output/antrean_kontinu_komparasi.png", width=Inches(5.8))
    
    insert_caption_before(target_p, "Gambar 4.8 Grafik Perbandingan Panjang Antrean Kontinu (Fixed-Time vs Adaptif Fuzzy)")
    
    insert_body_paragraph_before(target_p,
        "Kesuksesan efisiensi lalu lintas ini didorong secara langsung oleh berjalannya fungsi **Green Extension** dan "
        "**Green Cut** secara reaktif. Mekanisme *Green Extension* mendeteksi lajur pendekat yang masih mengalirkan antrean padat "
        "dan memperpanjang fase hijau secara bertahap hingga kendaraan habis atau menyentuh batas maksimum. Sebaliknya, "
        "mekanisme *Green Cut* bekerja memotong durasi hijau seketika ketika detektor YOLOv11s membaca lajur pendekat telah kosong, "
        "sehingga sisa detik hijau dapat langsung dialokasikan ke lajur pendekat lain yang sedang mengalami antrean."
    )
    
    insert_body_paragraph_before(target_p,
        "Selain itu, pembatasan durasi merah maksimum (**Max Red-Light Duration**) yang tetap terjaga stabil pada kapasitas **38,0 detik** "
        "bertindak sebagai batas konstrain (*boundary constraint*) yang aman. Batasan ini menjamin lajur pendekat lain yang sepi "
        "tidak akan mengalami penundaan tanpa batas (*starvation*) akibat lajur pendekat utama terus menerus memperpanjang fase hijaunya. "
        "Hasil akhir menunjukkan perbaikan indeks kepadatan persimpangan (**Density Index**) secara total sebesar **60,98% lebih lancar** "
        "(turun dari kepadatan 20,5% ke 8,0%), membuktikan sistem adaptif ini bekerja secara optimal untuk mengurai kemacetan perkotaan."
    )
    
    print("[+] Successfully inserted all new 4.3 sections and Gambar 4.8.")
    doc.save(doc_path)
    print("[OK] Rebuilding section 4.3 complete!")

if __name__ == "__main__":
    main()

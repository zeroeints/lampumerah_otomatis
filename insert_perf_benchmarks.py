import docx
from docx.shared import Inches, Pt, Cm
import os

def insert_body_paragraph_after(paragraph, text):
    p = paragraph.insert_paragraph_after()
    p.alignment = 3  # Justified
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def main():
    doc_path = "2215354055_Mohammad Filla Firdaus bab 1^J2^03.docx"
    doc = docx.Document(doc_path)
    print(f"Loaded document with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables.")
    
    # Let's find Tabel 4.6
    # In python-docx, tables are separate from paragraphs, but we can find where they are in the XML,
    # or we can find the paragraph that contains the caption 'Tabel 4.6' or is just before the table.
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "Tabel 4.6" in p.text:
            target_idx = i
            break
            
    if target_idx is None:
        print("[!] Could not find 'Tabel 4.6' paragraph caption.")
        return
        
    print(f"[+] Found 'Tabel 4.6' caption at paragraph index: {target_idx}")
    
    # We want to insert the narrative and the images AFTER the table.
    # The table is in the document structure. How to find the paragraph after the table?
    # In python-docx, if we find the paragraph containing 'Hasil pengujian integrasi komputasional membuktikan secara empiris', 
    # which is right after Table 4.6, we can insert our new elements before it.
    
    insert_before_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "Hasil pengujian integrasi komputasional membuktikan secara empiris" in p.text:
            insert_before_idx = i
            break
            
    if insert_before_idx is None:
        # Fallback to paragraph 347/348
        insert_before_idx = target_idx + 1
        
    print(f"[+] Will insert before paragraph index: {insert_before_idx}")
    target_p = doc.paragraphs[insert_before_idx]
    
    # Insert heading or narrative paragraph
    p_intro = target_p.insert_paragraph_before()
    p_intro.alignment = 3  # Justified
    p_intro.paragraph_format.first_line_indent = Cm(1.0)
    p_intro.paragraph_format.line_spacing = 1.5
    p_intro.paragraph_format.space_before = Pt(6)
    p_intro.paragraph_format.space_after = Pt(6)
    
    run_intro = p_intro.add_run(
        "Selain parameter efisiensi rekayasa lalu lintas simpang secara makro, pengujian tingkat sistem ini juga "
        "mengevaluasi kinerja komputasi dan stabilitas integrasi perangkat keras. Pengukuran parameter kinerja kualitatif "
        "ini meliputi latensi inferensi deteksi YOLOv11s, latensi keputusan pengontrol fuzzy, penggunaan memori RAM, dan "
        "stabilitas transmisi data protokol TraCI. Rincian hasil pengujian parameter operasional tersebut dipaparkan sebagai berikut:"
    )
    run_intro.font.name = 'Times New Roman'
    run_intro.font.size = Pt(12)
    
    # Add bullet points
    bullets = [
        ("Latensi Inferensi YOLOv11s: ", "Kecepatan pemrosesan frame citra visual model Small pada GPU Tesla T4 mencatatkan waktu rata-rata sebesar 9.7 ms per frame. Durasi komputasi ini setara dengan kecepatan penyegaran 103.1 FPS, memenuhi target konstrain operasional waktu nyata (< 100 ms)."),
        ("Latensi Keputusan Pengontrol Fuzzy: ", "Waktu penalaran matematis dari modul fuzzy_controller_opt.py dalam menyelesaikan satu siklus defuzzifikasi centroid adalah < 1.0 ms per langkah (dan rata-rata 4.3 ms pada pengujian sistem terintegrasi penuh), menjamin keputusan hijau diproduksi secara instan tanpa memicu desinkronisasi sirkuit clock simulator."),
        ("Alokasi Kapasitas Memori (RAM): ", "Konsumsi memori RAM total saat dashboard Streamlit, backend, dan sirkuit SUMO GUI beroperasi bersamaan stabil pada kapasitas ~2.73 GB s/d 2.8 GB, memenuhi target penalti aman penghematan hardware lokal (< 4 GB)."),
        ("Stabilitas Transmisi Protokol TraCI: ", "Komunikasi data dua arah antara Python dan kernel simulator mikro-trafik berjalan secara kontinu dan lancar tanpa mencatatkan adanya error pemutusan koneksi sepihak (0 errors / timeouts).")
    ]
    
    for title, desc in bullets:
        p_bullet = target_p.insert_paragraph_before()
        p_bullet.alignment = 3  # Justified
        p_bullet.paragraph_format.first_line_indent = Cm(1.0)
        p_bullet.paragraph_format.line_spacing = 1.5
        p_bullet.paragraph_format.space_before = Pt(0)
        p_bullet.paragraph_format.space_after = Pt(4)
        
        r_title = p_bullet.add_run(title)
        r_title.bold = True
        r_title.font.name = 'Times New Roman'
        r_title.font.size = Pt(12)
        
        r_desc = p_bullet.add_run(desc)
        r_desc.font.name = 'Times New Roman'
        r_desc.font.size = Pt(12)

    # Insert Image 4.6
    p_img1 = target_p.insert_paragraph_before()
    p_img1.alignment = 1  # Centered
    p_img1.paragraph_format.space_before = Pt(12)
    p_img1.paragraph_format.space_after = Pt(4)
    run_img1 = p_img1.add_run()
    run_img1.add_picture("output/latensi_komputasi_sistem.png", width=Inches(5.5))
    
    p_cap1 = target_p.insert_paragraph_before()
    p_cap1.alignment = 1  # Centered
    p_cap1.paragraph_format.space_before = Pt(4)
    p_cap1.paragraph_format.space_after = Pt(12)
    run_cap1 = p_cap1.add_run("Gambar 4.6 Grafik Perbandingan Latensi Pemrosesan Komputasi Sistem")
    run_cap1.bold = True
    run_cap1.font.name = 'Times New Roman'
    run_cap1.font.size = Pt(11)
    
    # Insert Image 4.7
    p_img2 = target_p.insert_paragraph_before()
    p_img2.alignment = 1  # Centered
    p_img2.paragraph_format.space_before = Pt(12)
    p_img2.paragraph_format.space_after = Pt(4)
    run_img2 = p_img2.add_run()
    run_img2.add_picture("output/penggunaan_memori_ram.png", width=Inches(5.5))
    
    p_cap2 = target_p.insert_paragraph_before()
    p_cap2.alignment = 1  # Centered
    p_cap2.paragraph_format.space_before = Pt(4)
    p_cap2.paragraph_format.space_after = Pt(12)
    run_cap2 = p_cap2.add_run("Gambar 4.7 Grafik Distribusi Alokasi Memori RAM Saat Sistem Beroperasi")
    run_cap2.bold = True
    run_cap2.font.name = 'Times New Roman'
    run_cap2.font.size = Pt(11)
    
    print("[+] Successfully inserted narrative benchmarks and two images before paragraph.")
    doc.save(doc_path)
    print("[OK] Word document saved successfully!")

if __name__ == "__main__":
    main()
